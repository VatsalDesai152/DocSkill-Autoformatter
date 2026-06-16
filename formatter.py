"""
formatter.py - Professional document formatting module.

Applies Times New Roman font with correct sizes:
  Title:    18pt, Bold
  Headings: 12pt, Bold
  Body:     11pt

CRITICAL: This module is careful NOT to destroy OMML equation XML
or any other special paragraph content when reformatting.
"""

from copy import deepcopy
from lxml import etree
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from equation_converter import convert_equations_in_doc


# OMML namespace used by Word for equations
OMML_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'


class DocumentFormatter:
    """Applies professional formatting to a DOCX document."""

    def __init__(self, doc_path):
        self.doc_path = doc_path
        self.doc = Document(doc_path)

    def format(self, output_path):
        """Full formatting pipeline."""
        # Step 1: Convert any LaTeX math ($...$, $$...$$) to native OMML
        eq_count = convert_equations_in_doc(self.doc)
        if eq_count > 0:
            print(f"  Converted {eq_count} LaTeX equation(s) to native Word format.")

        # Step 2: Apply styles and formatting
        self._apply_styles()
        self._format_paragraph_runs()
        self._format_table_runs()
        self._number_captions()
        self.doc.save(output_path)
        print(f"  Formatted document saved: {output_path}")

    # ------------------------------------------------------------------
    # Style-level formatting (affects defaults)
    # ------------------------------------------------------------------

    def _apply_styles(self):
        """Set font defaults on built-in styles."""
        for style in self.doc.styles:
            if style.type not in (WD_STYLE_TYPE.PARAGRAPH, WD_STYLE_TYPE.CHARACTER):
                continue
            if not hasattr(style, 'font'):
                continue

            if style.name == 'Title':
                self._set_font(style.font, size=18, bold=True)
            elif style.name.startswith('Heading'):
                self._set_font(style.font, size=12, bold=True)
            else:
                self._set_font(style.font, size=11)

    # ------------------------------------------------------------------
    # Run-level formatting (paragraph body text)
    # ------------------------------------------------------------------

    def _format_paragraph_runs(self):
        """Format runs in body paragraphs, SKIPPING any paragraph that
        contains OMML math elements so equations are never corrupted."""
        for para in self.doc.paragraphs:
            # --- GUARD: skip paragraphs with equations entirely ---
            if self._has_math(para):
                continue

            is_title = para.style.name == 'Title'
            is_heading = para.style.name.startswith('Heading')

            for run in para.runs:
                if is_title:
                    self._set_font(run.font, size=18, bold=True)
                elif is_heading:
                    self._set_font(run.font, size=12, bold=True)
                else:
                    self._set_font(run.font, size=11)

    def _format_table_runs(self):
        """Format text inside table cells."""
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if self._has_math(para):
                            continue
                        for run in para.runs:
                            self._set_font(run.font, size=11)

    # ------------------------------------------------------------------
    # Caption numbering
    # ------------------------------------------------------------------

    def _number_captions(self):
        """Standardise figure/table captions to 'Figure N: ...' / 'Table N: ...'

        IMPORTANT: We only modify the *text content* of runs, we never
        overwrite the paragraph element itself (p.text = ...) because
        that destroys the underlying XML including any images or equations.
        """
        fig_num = 1
        tbl_num = 1

        for para in self.doc.paragraphs:
            if self._has_math(para):
                continue

            text = para.text.strip()
            if not text:
                continue

            low = text.lower()

            if low.startswith(('figure', 'fig.', 'fig ')):
                detail = text.split(':', 1)[1].strip() if ':' in text else ''
                new_text = f"Figure {fig_num}: {detail}"
                self._rewrite_runs_text(para, new_text)
                fig_num += 1

            elif low.startswith('table') and not low.startswith('table of'):
                detail = text.split(':', 1)[1].strip() if ':' in text else ''
                new_text = f"Table {tbl_num}: {detail}"
                self._rewrite_runs_text(para, new_text)
                tbl_num += 1

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _set_font(font, size, bold=None):
        """Apply Times New Roman at the given point size."""
        font.name = 'Times New Roman'
        font.size = Pt(size)
        if bold is not None:
            font.bold = bold

    @staticmethod
    def _has_math(para):
        """Return True if the paragraph XML contains OMML math elements."""
        return para._element.findall(qn('m:oMath')) or \
               para._element.findall(qn('m:oMathPara'))

    @staticmethod
    def _rewrite_runs_text(para, new_text):
        """Safely rewrite paragraph text across existing runs
        without destroying the underlying XML structure."""
        runs = para.runs
        if not runs:
            return
        # Put all new text in the first run, clear the rest
        runs[0].text = new_text
        for run in runs[1:]:
            run.text = ''
