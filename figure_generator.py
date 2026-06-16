"""
figure_generator.py - Automatic figure generation from document tables.

Reads tables from a DOCX, identifies ones with numeric data,
generates matplotlib bar charts, and inserts them back into the document.
"""

import os
import pandas as pd
import matplotlib
matplotlib.use('Agg')  # Non-interactive backend
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


class FigureGenerator:
    """Generates figures from numeric tables in a DOCX and inserts them."""

    def __init__(self, doc_path):
        self.doc_path = doc_path
        self.doc = Document(doc_path)

    def generate_and_insert_figures(self, output_dir='output_figures'):
        """Main entry: iterate tables, generate plots, insert into doc."""
        os.makedirs(output_dir, exist_ok=True)

        generated = 0
        for idx, table in enumerate(self.doc.tables):
            df = self._table_to_dataframe(table)
            if df is None:
                continue

            fig_path = os.path.join(output_dir, f'generated_fig_{idx + 1}.png')
            if self._create_chart(df, idx + 1, fig_path):
                self._insert_after_table(table, fig_path, idx + 1)
                generated += 1

        if generated > 0:
            self.doc.save(self.doc_path)
            print(f"  Generated {generated} figure(s) and inserted into document.")
        else:
            print("  No numeric tables found for figure generation.")

    # ------------------------------------------------------------------
    # Table parsing
    # ------------------------------------------------------------------

    @staticmethod
    def _table_to_dataframe(table):
        """Convert a docx table to a pandas DataFrame.
        Returns None if table has no usable numeric data."""
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])

        if len(rows) < 2 or len(rows[0]) < 1:
            return None

        df = pd.DataFrame(rows[1:], columns=rows[0])

        # Check for at least one numeric column
        numeric = df.apply(pd.to_numeric, errors='coerce')
        has_numbers = numeric.dropna(axis=1, how='all')
        if has_numbers.empty:
            return None

        return df

    # ------------------------------------------------------------------
    # Chart creation
    # ------------------------------------------------------------------

    @staticmethod
    def _create_chart(df, table_num, save_path):
        """Create a bar chart from numeric columns. Returns True on success."""
        try:
            numeric = df.apply(pd.to_numeric, errors='coerce')
            valid = numeric.dropna(axis=1, how='all')

            fig, ax = plt.subplots(figsize=(6, 4))

            # Use first column as labels if it's non-numeric
            first_col = df.columns[0]
            if first_col not in valid.columns:
                valid.index = df[first_col].values

            valid.plot(kind='bar', ax=ax)

            ax.set_title(f'Generated Figure for Table {table_num}',
                         fontname='Times New Roman', fontsize=12, fontweight='bold')
            ax.tick_params(axis='x', rotation=45)
            for label in ax.get_xticklabels() + ax.get_yticklabels():
                label.set_fontname('Times New Roman')
                label.set_fontsize(11)

            plt.tight_layout()
            plt.savefig(save_path, dpi=300)
            plt.close(fig)
            return True

        except Exception as e:
            print(f"  Could not generate figure for table {table_num}: {e}")
            plt.close('all')
            return False

    # ------------------------------------------------------------------
    # Insertion into DOCX
    # ------------------------------------------------------------------

    def _insert_after_table(self, table, image_path, table_num):
        """Insert an image + caption paragraph right after the given table."""
        # Image paragraph
        p_img_el = OxmlElement('w:p')
        table._element.addnext(p_img_el)
        p_img = Paragraph(p_img_el, self.doc._body)
        run = p_img.add_run()
        run.add_picture(image_path, width=Inches(5.0))

        # Caption paragraph
        p_cap_el = OxmlElement('w:p')
        p_img_el.addnext(p_cap_el)
        p_cap = Paragraph(p_cap_el, self.doc._body)
        cap_run = p_cap.add_run(f'Figure: Automatically generated plot from Table {table_num}')
        cap_run.font.name = 'Times New Roman'
        cap_run.font.size = Pt(11)
